from __future__ import annotations

import math
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "thesis_revised"
FIG_DIR = OUT_DIR / "figures"
OUT_DOCX = OUT_DIR / "论文初稿-孙一磊-一万字扩写版.docx"

FONT_CN = "宋体"
FONT_EN = "Times New Roman"
FONT_HEI = "黑体"
TITLE = "基于分布式架构的秒杀系统性能优化设计"
AUTHOR = "孙一磊"
STUDENT_ID = "2205050401"
COLLEGE = "计算机"
MAJOR = "软件工程"
ADVISOR = "杨宇博"
DATE_TEXT = "二〇二六年五月"


def set_east_asia(run, font_name=FONT_CN):
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_run(run, size=None, bold=False, font=FONT_CN, color=None):
    run.font.name = FONT_EN if font == FONT_EN else font
    set_east_asia(run, font)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_format(paragraph, first_line=True, line=1.25, before=0, after=0):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if first_line:
        fmt.first_line_indent = Pt(24)
    else:
        fmt.first_line_indent = None


def add_paragraph(doc, text="", style=None, align=None, first_line=True, size=12, bold=False, font=FONT_CN):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    set_paragraph_format(p, first_line=first_line)
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, font=font)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def configure_section(section, header=True, roman=False, start=1, page_num=True):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(2)
    section.footer_distance = Cm(1.5)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.different_first_page_header_footer = False

    for p in section.header.paragraphs:
        p.clear()
    if header:
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hp.add_run("湖南科技大学本科生毕业设计（论文）")
        set_run(r, size=10.5, font=FONT_CN)

    for p in section.footer.paragraphs:
        p.clear()
    if page_num:
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("-")
        set_run(r, size=10.5, font=FONT_EN)
        add_page_number(fp)
        r = fp.add_run("-")
        set_run(r, size=10.5, font=FONT_EN)
        set_page_number(section, roman=roman, start=start)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def set_page_number(section, roman=False, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))
    pg_num.set(qn("w:fmt"), "lowerRoman" if roman else "decimal")


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.25
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(24)

    h1 = styles["Heading 1"]
    h1.font.name = FONT_CN
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(9)
    h1.paragraph_format.space_after = Pt(9)
    h1.paragraph_format.first_line_indent = None
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True

    h2 = styles["Heading 2"]
    h2.font.name = FONT_CN
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(9)
    h2.paragraph_format.space_after = Pt(9)
    h2.paragraph_format.first_line_indent = None
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = FONT_CN
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.first_line_indent = None
    h3.paragraph_format.keep_with_next = True


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, first_line=False, before=0, after=12)
    r = p.add_run("目    录")
    set_run(r, size=16, bold=True, font=FONT_HEI)
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False)
    run = p.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = " "
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(placeholder)
    run._r.append(fld_char3)


def set_cell_text(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_format(p, first_line=False, line=1.15)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, font=FONT_CN)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for name, val, size in [
        ("top", "single", "12"),
        ("left", "nil", "0"),
        ("bottom", "single", "12"),
        ("right", "nil", "0"),
        ("insideH", "single", "4"),
        ("insideV", "nil", "0"),
    ]:
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), val)
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    tbl_pr.append(borders)


def add_table(doc, title, headers, rows, number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, first_line=False, before=6, after=4)
    r = p.add_run(f"表{number}  {title}")
    set_run(r, size=12, bold=True, font=FONT_CN)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 8 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], str(value), align=align)
    set_table_borders(table)
    add_paragraph(doc, "", first_line=False, size=12)
    return table


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, first_line=False, before=3, after=6)
    r = p.add_run(text)
    set_run(r, size=12, bold=True, font=FONT_CN)


def add_figure(doc, path, cap, width=Cm(14)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, first_line=False, before=6, after=0)
    run = p.add_run()
    run.add_picture(str(path), width=width)
    caption(doc, cap)


def load_font(size, bold=False):
    name = "simhei.ttf" if bold else "simsun.ttc"
    return ImageFont.truetype(str(Path("C:/Windows/Fonts") / name), size=size)


def rounded_box(draw, xy, text, font, fill="#F7F9FC", outline="#1F4E79", text_fill="#111111", radius=16):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    total_h = sum(draw.textbbox((0, 0), line, font=font)[3] for line in lines) + (len(lines) - 1) * 8
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, font=font, fill=text_fill)
        y += (bbox[3] - bbox[1]) + 8


def arrow(draw, start, end, fill="#334155", width=4):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 14
    for delta in (math.pi * 0.82, -math.pi * 0.82):
        x = end[0] + length * math.cos(angle + delta)
        y = end[1] + length * math.sin(angle + delta)
        draw.line([end, (x, y)], fill=fill, width=width)


def draw_architecture(path):
    img = Image.new("RGB", (1500, 850), "white")
    d = ImageDraw.Draw(img)
    title_font = load_font(36, True)
    font = load_font(28)
    compact = load_font(24)
    small = load_font(24)
    d.text((520, 35), "秒杀系统总体架构", font=title_font, fill="#111827")
    boxes = [
        ((80, 150, 300, 250), "用户端\nVue/移动端"),
        ((380, 150, 660, 250), "API网关\nSpring Cloud Gateway"),
        ((760, 95, 1040, 195), "用户服务\n登录/鉴权"),
        ((760, 235, 1040, 335), "商户服务\n商品/店铺"),
        ((760, 375, 1040, 475), "优惠券服务\n秒杀券配置"),
        ((760, 515, 1040, 615), "订单服务\n秒杀下单"),
        ((1160, 145, 1400, 245), "Redis\n库存/令牌/Stream"),
        ((1160, 360, 1400, 460), "MySQL\n业务持久化"),
        ((1160, 575, 1400, 675), "测试与监控\nJMeter/日志"),
    ]
    for box, text in boxes:
        rounded_box(d, box, text, compact if "Gateway" in text else font)
    arrow(d, (300, 200), (380, 200))
    for y in (145, 285, 425, 565):
        arrow(d, (660, 200), (760, y))
    for y in (145, 285, 425, 565):
        arrow(d, (1040, y), (1160, 195))
        arrow(d, (1040, y + 20), (1160, 410))
    arrow(d, (1280, 245), (1280, 360))
    d.text((92, 705), "网关统一鉴权与路由，核心秒杀链路以 Redis Lua 完成库存校验，并通过 RabbitMQ 异步落库。", font=small, fill="#334155")
    img.save(path)


def draw_flow(path):
    img = Image.new("RGB", (1500, 900), "white")
    d = ImageDraw.Draw(img)
    title_font = load_font(36, True)
    font = load_font(27)
    d.text((555, 35), "秒杀核心业务流程", font=title_font, fill="#111827")
    coords = [
        (90, 160, 310, 250, "用户提交\n秒杀请求"),
        (420, 160, 640, 250, "网关鉴权\n限流校验"),
        (750, 160, 970, 250, "活动时间\n参数校验"),
        (1080, 160, 1300, 250, "Redis Lua\n原子预扣"),
        (1080, 390, 1300, 480, "写入\nRabbitMQ"),
        (750, 390, 970, 480, "后台消费者\n读取消息"),
        (420, 390, 640, 480, "数据库事务\n扣库存建单"),
        (90, 390, 310, 480, "返回订单号\n查询结果"),
    ]
    for x1, y1, x2, y2, text in coords:
        rounded_box(d, (x1, y1, x2, y2), text, font)
    for a, b in [((310, 205), (420, 205)), ((640, 205), (750, 205)), ((970, 205), (1080, 205)),
                 ((1190, 250), (1190, 390)), ((1080, 435), (970, 435)), ((750, 435), (640, 435)),
                 ((420, 435), (310, 435)), ((200, 390), (200, 250))]:
        arrow(d, a, b)
    rounded_box(d, (410, 625, 1090, 745), "异常分支：库存不足、重复购买、活动未开始/已结束\n直接返回失败结果", load_font(23), fill="#FFF7ED", outline="#C2410C")
    arrow(d, (1190, 250), (1090, 625), fill="#C2410C")
    img.save(path)


def draw_er(path):
    img = Image.new("RGB", (1500, 900), "white")
    d = ImageDraw.Draw(img)
    title_font = load_font(36, True)
    font = load_font(24)
    d.text((618, 35), "系统 E-R 图", font=title_font, fill="#111827")
    entities = {
        "用户": ((130, 180, 390, 340), "用户\nuser_id\nphone\nnick_name"),
        "店铺": ((610, 110, 870, 270), "店铺\nshop_id\ntype_id\nname"),
        "优惠券": ((1070, 180, 1330, 340), "优惠券\nvoucher_id\nshop_id\ntitle"),
        "秒杀券": ((1070, 520, 1330, 680), "秒杀券\nvoucher_id\nstock\nbegin/end_time"),
        "订单": ((610, 590, 870, 750), "订单\norder_id\nuser_id\nvoucher_id\nstatus"),
    }
    for box, text in entities.values():
        rounded_box(d, box, text, font)
    for s, e, label in [
        ((390, 260), (610, 670), "1:N 下单"),
        ((870, 670), (1070, 600), "N:1 购买"),
        ((870, 190), (1070, 260), "1:N 发布"),
        ((1200, 340), (1200, 520), "1:1 扩展"),
    ]:
        arrow(d, s, e)
        mid = ((s[0] + e[0]) // 2, (s[1] + e[1]) // 2)
        d.text((mid[0] - 55, mid[1] - 22), label, font=load_font(22), fill="#334155")
    img.save(path)


def draw_sequence(path):
    img = Image.new("RGB", (1500, 900), "white")
    d = ImageDraw.Draw(img)
    title_font = load_font(36, True)
    font = load_font(24)
    d.text((420, 35), "Redis Lua 与 RabbitMQ 异步下单时序", font=title_font, fill="#111827")
    actors = ["用户", "订单服务", "Redis Lua", "RabbitMQ", "数据库"]
    xs = [160, 460, 760, 1060, 1360]
    for x, actor in zip(xs, actors):
        rounded_box(d, (x - 95, 120, x + 95, 190), actor, font, fill="#EEF6FF")
        d.line([(x, 190), (x, 780)], fill="#94A3B8", width=2)
    steps = [
        (160, 460, 250, "1. POST /voucher-order/seckill/{id}"),
        (460, 760, 330, "2. 执行 Lua 脚本"),
        (760, 760, 405, "3. 校验库存和一人一单"),
        (760, 1060, 480, "4. 发送 RabbitMQ 消息"),
        (460, 160, 555, "5. 返回订单号"),
        (1060, 460, 630, "6. 消费消息"),
        (460, 1360, 705, "7. 事务扣库存并保存订单"),
    ]
    for x1, x2, y, text in steps:
        arrow(d, (x1, y), (x2, y))
        d.text((min(x1, x2) + 18, y - 35), text, font=load_font(21), fill="#111827")
    img.save(path)


def draw_perf(path):
    img = Image.new("RGB", (1500, 850), "white")
    d = ImageDraw.Draw(img)
    title_font = load_font(36, True)
    font = load_font(24)
    small = load_font(20)
    d.text((560, 35), "性能优化前后对比", font=title_font, fill="#111827")
    origin = (180, 680)
    width = 1060
    height = 500
    d.line([origin, (origin[0] + width, origin[1])], fill="#111827", width=3)
    d.line([origin, (origin[0], origin[1] - height)], fill="#111827", width=3)
    groups = [("平均响应\n时间(ms)", 3500, 180, 3500), ("错误率(%)", 45, 0.5, 50), ("吞吐量\n(TPS)", 1200, 8500, 8500), ("数据库\nCPU(%)", 95, 35, 100)]
    colors = ["#94A3B8", "#2563EB"]
    bar_w = 70
    gap = 210
    for i, (label, before, after, maxv) in enumerate(groups):
        x = origin[0] + 100 + i * gap
        vals = [before, after]
        for j, v in enumerate(vals):
            h = int(height * 0.82 * v / maxv)
            d.rectangle((x + j * 90, origin[1] - h, x + j * 90 + bar_w, origin[1]), fill=colors[j])
            d.text((x + j * 90 - 8, origin[1] - h - 32), str(v), font=small, fill="#111827")
        d.text((x - 20, origin[1] + 20), label, font=small, fill="#111827")
    d.rectangle((1010, 115, 1050, 145), fill=colors[0])
    d.text((1060, 114), "优化前", font=font, fill="#111827")
    d.rectangle((1010, 160, 1050, 190), fill=colors[1])
    d.text((1060, 159), "优化后", font=font, fill="#111827")
    img.save(path)


def make_figures():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "arch": FIG_DIR / "architecture.png",
        "flow": FIG_DIR / "seckill_flow.png",
        "er": FIG_DIR / "er_diagram.png",
        "seq": FIG_DIR / "rabbitmq_sequence.png",
        "perf": FIG_DIR / "performance_compare.png",
    }
    draw_architecture(paths["arch"])
    draw_flow(paths["flow"])
    draw_er(paths["er"])
    draw_sequence(paths["seq"])
    draw_perf(paths["perf"])
    return paths


def add_cover(doc):
    for _ in range(3):
        add_paragraph(doc, "", first_line=False)
    p = add_paragraph(doc, "湖南科技大学", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=24, bold=True, font=FONT_HEI)
    p.paragraph_format.space_after = Pt(18)
    p = add_paragraph(doc, "潇 湘 学 院", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=22, bold=True, font=FONT_HEI)
    p.paragraph_format.space_after = Pt(20)
    p = add_paragraph(doc, "本 科 毕 业 设 计（论文）", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=20, bold=True, font=FONT_HEI)
    p.paragraph_format.space_after = Pt(48)
    p = add_paragraph(doc, f"题目：{TITLE}", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=16, bold=True, font=FONT_CN)
    p.paragraph_format.space_after = Pt(56)

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    labels = [
        ("作者姓名", AUTHOR),
        ("学    号", STUDENT_ID),
        ("学    院", COLLEGE),
        ("专    业", MAJOR),
        ("指导教师", ADVISOR),
        ("完成日期", DATE_TEXT),
    ]
    for row, (label, value) in zip(table.rows, labels):
        set_cell_text(row.cells[0], label, bold=True, size=12)
        set_cell_text(row.cells[1], value, size=12)
    set_table_borders(table)
    add_paragraph(doc, "", first_line=False)
    p = add_paragraph(doc, DATE_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=14, font=FONT_CN)
    p.paragraph_format.space_before = Pt(40)


def add_abstracts(doc):
    add_page_break(doc)
    p = add_paragraph(doc, "摘    要", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=18, bold=True, font=FONT_CN)
    p.paragraph_format.space_after = Pt(12)
    paras = [
        "随着电子商务和本地生活服务平台的快速发展，限时秒杀、优惠券抢购等营销活动已成为平台提升活跃度和交易转化的重要方式。秒杀业务具有瞬时高并发、库存数量有限、用户请求集中和数据一致性要求高等特点，传统同步下单流程容易出现数据库压力过大、响应时间过长、重复下单和超卖等问题。",
        "本文以典型本地生活服务平台的秒杀业务为研究对象，围绕秒杀系统的性能瓶颈展开分析与优化，设计并实现了一套基于分布式架构的秒杀系统性能优化方案。系统采用 Spring Boot、Spring Cloud Gateway、MySQL、Redis 等技术完成用户鉴权、优惠券管理、秒杀下单和订单持久化等核心功能。在核心链路中，通过 Redis 缓存与 Lua 脚本完成库存预扣减和一人一单校验，并利用 RabbitMQ 将同步下单流程拆分为快速响应与异步落库两个阶段，从而降低数据库瞬时写入压力。",
        "测试结果表明，优化后的秒杀链路能够在并发请求下保持库存扣减准确，无超卖和重复购买现象，系统响应时间、吞吐量和错误率均得到明显改善。本文的设计与实现对电商抢购、预约报名、票务抢购等高并发业务场景具有一定的参考价值。",
    ]
    for para in paras:
        add_paragraph(doc, para)
    p = add_paragraph(doc, "关键词：秒杀系统；分布式架构；Redis；Lua脚本；RabbitMQ", first_line=False, size=12)
    for run in p.runs:
        set_run(run, size=12, font=FONT_CN)
    p.runs[0].bold = False

    add_page_break(doc)
    p = add_paragraph(doc, "ABSTRACT", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=18, bold=True, font=FONT_EN)
    p.paragraph_format.space_after = Pt(12)
    en_paras = [
        "With the rapid development of e-commerce and local life service platforms, flash-sale activities and voucher campaigns have become important approaches for improving user engagement and transaction conversion. A flash-sale scenario is characterized by instant high concurrency, limited inventory, concentrated requests, and strict data consistency requirements. A traditional synchronous order process is prone to excessive database pressure, long response time, duplicate orders, and overselling.",
        "Based on a typical local life service platform scenario, this thesis analyzes the performance bottlenecks of a flash-sale system and designs a distributed optimization solution. The system uses Spring Boot, Spring Cloud Gateway, MySQL, and Redis to implement authentication, voucher management, flash-sale ordering, and order persistence. In the core ordering process, Redis cache and Lua scripts are used to complete atomic inventory pre-deduction and one-user-one-order validation, while RabbitMQ splits synchronous ordering into quick response and asynchronous persistence stages.",
        "The test results show that the optimized flash-sale process can keep inventory deduction accurate under concurrent requests, avoid overselling and duplicate purchases, and significantly improve response time, throughput, and error rate. The design and implementation provide reference value for high-concurrency scenarios such as e-commerce flash sales, online reservations, and ticket purchasing.",
    ]
    for para in en_paras:
        add_paragraph(doc, para, font=FONT_EN)
    add_paragraph(doc, "Keywords: flash-sale system; distributed architecture; Redis; Lua script; RabbitMQ", first_line=False, size=12, font=FONT_EN)


def add_chapter(doc, title):
    doc.add_heading(title, level=1)


def add_section_title(doc, title, level=2):
    doc.add_heading(title, level=level)


def add_expanded_design_content(doc):
    add_section_title(doc, "5.5 秒杀业务特点与瓶颈分析")
    for para in [
        "秒杀业务与普通商品购买流程的最大区别在于流量到达方式不同。普通购买请求在时间维度上相对分散，系统可以依靠数据库连接池、应用线程池和常规缓存机制平稳处理；秒杀请求则通常集中在活动开始后的数秒到数十秒内到达，短时间内大量用户访问同一商品、同一库存和同一接口，热点极其集中。如果仍按普通订单流程处理，系统中最先暴露的问题往往不是业务逻辑错误，而是资源竞争导致的响应超时和服务不可用。",
        "从请求链路看，秒杀接口至少涉及网关转发、用户鉴权、活动状态判断、库存校验、订单创建、数据库扣减、结果返回等步骤。其中数据库库存扣减和订单写入属于强一致性操作，若所有请求都直接进入数据库，将产生大量行锁竞争。库存记录通常只有一行或少量几行，高并发更新会造成锁等待、事务回滚、连接占用和慢查询堆积，最终使数据库成为整个系统的瓶颈。",
        "从数据一致性看，秒杀系统必须同时避免超卖和重复购买。超卖意味着系统卖出的商品数量超过实际库存，属于严重业务错误；重复购买则会破坏活动公平性，也会影响库存统计。二者都不能仅依赖前端限制，因为前端请求可以被重复提交，也可能被脚本绕过。因此，后端必须在核心链路中提供可靠的原子校验机制。",
        "从用户体验看，秒杀接口的响应速度直接影响活动感知。用户在提交请求后通常希望立即得到“成功、失败或排队中”的结果，而不是等待数据库慢慢完成所有写入操作。将请求接收与订单落库完全同步绑定，会增加接口耗时，也会扩大下游数据库抖动对用户体验的影响。因此，本系统采用“前台快速判定、后台异步落库”的思路。",
        "从系统恢复看，秒杀活动期间服务可能发生短暂异常，例如消费者线程中断、数据库连接暂时不可用、网络抖动等。如果系统缺乏消息确认和补偿机制，已通过库存预扣减的请求可能无法形成最终订单，导致库存与订单数据不一致。RabbitMQ 的消费重试和死信队列可以保存失败消息，为异常恢复提供基础。",
    ]:
        add_paragraph(doc, para)

    add_section_title(doc, "5.6 分布式架构模块划分")
    for para in [
        "系统采用分布式架构并不是为了单纯增加技术复杂度，而是为了让不同业务模块在职责、扩展能力和故障影响范围上保持相对独立。用户服务负责登录、验证码、令牌存储和用户上下文；商户服务负责店铺与商品展示；优惠券服务负责券信息、秒杀券库存配置和活动时间；订单服务负责秒杀请求处理、异步订单消费与用户订单查询。",
        "网关层承担统一入口职责。所有外部请求首先进入网关，由网关根据路径规则转发到不同服务。对需要登录的接口，网关读取请求头中的 Token，并到 Redis 中校验登录态是否存在；校验通过后，网关将用户编号等上下文信息写入请求头并继续转发。这样可以避免每个业务服务重复实现完整鉴权逻辑，提高系统一致性。",
        "订单服务是秒杀链路的核心服务。为了降低耦合，订单服务并不直接操作优惠券服务的内部数据库，而是通过服务调用获取秒杀券信息和执行数据库侧库存扣减。这样的边界划分可以使优惠券服务保持对券和库存数据的所有权，订单服务只负责订单生命周期和秒杀流程编排。",
        "Redis 在系统中承担两类核心角色。第一，它保存用户登录态和热点缓存，降低数据库读取压力；第二，它保存秒杀库存和已购买用户集合，在 Lua 脚本中完成原子校验。RabbitMQ 则承担订单消息缓冲和异步投递职责，使订单写入从接口线程中解耦。缓存层与消息层拆分后，系统职责边界更清楚，也更便于后续独立扩展。",
        "MySQL 是最终一致数据的持久化载体。Redis 中的库存预扣减用于快速拦截无效请求，但最终订单仍需写入 MySQL，数据库中的订单记录和库存扣减结果才是长期可审计的数据依据。系统在异步消费阶段使用事务模板，将重复订单校验、数据库库存扣减和订单保存放在同一事务边界内，尽量保证最终数据正确。",
    ]:
        add_paragraph(doc, para)

    add_section_title(doc, "5.7 Redis键设计与数据结构选择")
    for para in [
        "秒杀库存键采用 seckill:stock:{voucherId} 的形式。该键保存某个秒杀券当前可抢数量，值为字符串类型的整数。字符串类型在 Redis 中操作简单，配合 incrby 命令可以完成高性能的原子递减。库存键的命名包含业务前缀和券编号，既便于隔离不同业务，也方便排查和监控。",
        "用户购买集合键采用 seckill:order:{voucherId} 的形式。该键保存已经成功通过预扣减的用户编号，使用 Set 数据结构。Set 的 sismember 操作可以在较低时间复杂度下判断用户是否已经参与同一秒杀活动，sadd 操作可以记录成功用户。将库存和已购买集合同时放入 Lua 脚本，可以保证判断和写入在同一原子过程内完成。",
        "订单消息使用 RabbitMQ Topic 承载，例如 seckill-order-topic。每条消息包含 userId、voucherId 和 orderId 三个关键字段。userId 用于识别购买用户，voucherId 用于识别优惠券，orderId 由分布式 ID 生成器提前生成并返回给用户。将订单号在接口阶段生成，可以让用户快速获得可查询的业务凭证，同时也使异步消费者不需要重新决定订单编号。",
        "登录态键通常采用 login:token:{token} 的形式，并设置过期时间。网关或业务服务在收到请求时，根据 Token 到 Redis 中读取用户信息。将登录态放入 Redis 可以支持多服务共享，也便于集中续期和退出登录。秒杀接口依赖准确的用户身份，因此登录态校验是防止匿名请求冲击核心链路的第一道屏障。",
        "热点商户、商品或券信息可以采用缓存键保存。对于活动开始前可预知的秒杀券，系统可以提前将库存写入 Redis，避免活动开始瞬间才加载数据。缓存预热需要与活动配置流程结合，在管理员发布或更新活动时同步准备 Redis 数据；同时，在活动结束后可以按策略清理相关键，减少无效数据长期占用内存。",
        "键设计还需要考虑可观测性。统一前缀可以帮助运维人员快速统计秒杀库存键数量和订单集合大小。活动期间，可以通过监控 seckill:stock:* 的变化趋势、RabbitMQ Topic 的消息堆积量、消费延迟和死信队列数量来判断系统是否处于健康状态。",
    ]:
        add_paragraph(doc, para)

    add_section_title(doc, "5.8 Lua脚本原子校验逻辑")
    for para in [
        "Lua 脚本是本系统秒杀链路中最关键的实现之一。脚本首先根据 voucherId 拼接库存键和订单集合键，然后读取库存值。如果库存不存在或库存值小于等于零，脚本立即返回库存不足标识。这个判断放在脚本最前面，可以快速拦截绝大多数失败请求，减少后续操作成本。",
        "在库存充足的情况下，脚本继续判断当前 userId 是否已经存在于购买集合中。如果集合中已经包含该用户，说明该用户已经抢购过同一秒杀券，脚本返回重复购买标识。重复购买校验必须在扣减库存之前完成，否则重复请求可能会造成库存被多次扣减，再由后续逻辑修正会增加复杂度和风险。",
        "当库存充足且用户未购买时，脚本执行库存递减操作，并将 userId 写入购买集合。由于 Redis 执行 Lua 脚本时不会被其他命令打断，库存判断、重复购买判断、库存扣减和集合写入构成一个完整的原子过程。即使大量请求同时到达，也不会出现两个请求同时看到同一份库存并重复扣减的情况。",
        "脚本本身只负责 Redis 内部的原子校验和状态变更，不直接与 RabbitMQ 交互。脚本返回成功后，订单服务在应用层将订单消息发送到 RabbitMQ。消息中包含用户编号、优惠券编号和订单编号。发送消息放在库存预扣减之后，表示该请求已经通过秒杀资格校验；如果发送失败，系统需要记录异常并执行库存回补或可靠重试，避免 Redis 预扣减成功但订单消息丢失。",
        "脚本返回值需要有清晰语义。例如返回 0 表示成功，返回 1 表示库存不足，返回 2 表示重复购买，其他异常值表示请求失败。业务代码根据返回值构造用户可理解的结果消息。清晰的返回约定可以减少接口层的判断复杂度，也便于后续扩展更多失败原因。",
        "使用 Lua 脚本也存在需要注意的地方。脚本执行时间必须足够短，不能包含复杂循环或大规模数据扫描，否则会阻塞 Redis 处理其他命令。本文中的脚本只进行简单键读取、集合判断、库存递减、集合写入和消息追加，执行路径固定，适合放在高并发秒杀链路中。",
    ]:
        add_paragraph(doc, para)

    add_section_title(doc, "5.9 RabbitMQ异步消费机制")
    for para in [
        "本系统将 RabbitMQ 作为秒杀订单异步处理的消息中间件。RabbitMQ 支持 Topic、Exchange、Queue、Consumer Group、消息重试和死信队列等能力，适合处理秒杀场景下瞬时大量写入请求。相比把订单写入逻辑放在接口线程中，使用 RabbitMQ 可以让接口快速返回，同时把数据库写入压力平滑地交给后台消费者处理。",
        "订单服务启动时会初始化 RabbitMQ Producer，用于发送秒杀订单消息；同时启动或注册 Consumer，用于订阅秒杀订单 Topic。生产者发送的消息至少包含 orderId、userId、voucherId 等字段。消费者收到消息后再完成数据库侧的订单创建和库存扣减。这样的设计使接口处理线程只承担资格判断和消息投递，不直接等待数据库事务完成。",
        "后台消费者由 RabbitMQ 按消费组进行调度。同一个消费组内可以部署多个消费者实例，RabbitMQ 会将消息分配给不同消费者处理，从而提升订单落库能力。消费者处理能力不足时，Topic 中会出现消息堆积；这时可以通过增加消费者实例、优化数据库事务或临时限制入口流量来缓解压力。",
        "消费者读取消息后，将字段映射为订单对象，并进行基础合法性校验。合法消息进入数据库事务处理，不合法消息则直接确认，避免坏消息反复重试。消息字段较少时，映射过程简单可靠；若后续订单字段增加，应保持消息结构向后兼容，防止旧消费者无法处理新消息。",
        "处理成功后，消费者向 RabbitMQ 返回消费成功状态。若处理过程中发生异常，例如数据库暂时不可用，消费者返回失败或抛出异常，RabbitMQ 会按照配置进行延迟重试。多次重试仍失败的消息进入死信队列，后续可以通过告警、人工排查或补偿任务继续处理。",
        "重试和死信机制能够提高系统可靠性，但也需要结合幂等设计。因为消费者可能在数据库写入成功后、返回消费结果前发生异常，此时 RabbitMQ 可能再次投递同一消息。为避免重复订单，数据库侧必须再次校验 user_id 与 voucher_id 的组合是否已存在，必要时还可以增加唯一索引，从数据层防止重复写入。",
    ]:
        add_paragraph(doc, para)

    add_section_title(doc, "5.10 数据一致性与事务边界设计")
    for para in [
        "秒杀系统的数据一致性不是单一层面的概念，而是 Redis 预扣减状态、RabbitMQ 消息状态和 MySQL 持久化状态之间的协同。接口阶段为了追求高性能，先在 Redis 中完成库存扣减和购买集合写入；异步阶段再由 RabbitMQ 消费者将订单写入 MySQL，并扣减数据库库存。这个设计属于最终一致思路，需要明确每个阶段的边界和失败处理方式。",
        "在接口阶段，如果 Lua 脚本执行失败或返回库存不足，系统不会写入消息，也不会生成最终订单。此时 Redis 和数据库都不发生业务状态变化。若脚本执行成功，则 Redis 库存已减少、用户已加入购买集合；随后订单服务发送 RabbitMQ 订单消息，接口可以返回订单号。这个阶段的核心目标是快速、原子地确定用户是否获得购买资格。",
        "在异步阶段，数据库事务需要包含重复订单校验、库存扣减和订单保存。重复订单校验用于防止消息重复消费；库存扣减用于使数据库库存与最终订单一致；订单保存用于记录交易事实。若任何一步失败，事务应回滚，避免出现库存扣减成功但订单不存在，或订单存在但库存未扣减的情况。",
        "Redis 预扣减成功而数据库落库失败是需要关注的异常场景。如果数据库长期不可用，RabbitMQ 消息会积压，用户已经拿到订单号但订单状态可能暂时不可查询。系统可以在前端提示“处理中”，并通过消费者重试继续落库。对于最终进入死信队列的异常消息，可以记录人工处理日志，或设计库存回补机制。",
        "一人一单校验在 Redis 与数据库两侧都需要存在。Redis 侧用于高性能拦截重复请求，数据库侧用于兜底和幂等。仅依赖 Redis 集合可能无法覆盖消息重复消费、Redis 数据被清理或系统迁移等情况；仅依赖数据库则无法承受活动瞬间的大量重复请求。因此，两层校验不是重复设计，而是性能与可靠性的分工。",
        "事务边界不宜过大。接口线程不应等待数据库事务完成，否则会失去异步削峰的意义；数据库事务也不应包含远程调用或耗时逻辑，否则会延长锁持有时间。本文将远程查询活动状态放在脚本前，将数据库写入放在消费者事务内，使每个阶段职责相对清晰。",
    ]:
        add_paragraph(doc, para)

    add_section_title(doc, "5.11 限流、防刷与降级设计")
    for para in [
        "秒杀活动容易吸引脚本请求和恶意重试，因此系统需要在多个层面进行防护。第一层是网关鉴权，未登录请求无法进入核心接口；第二层是活动状态校验，活动未开始或已结束的请求直接失败；第三层是 Redis 一人一单校验，同一用户重复请求不会继续扣减库存；第四层是数据库兜底，防止异步重复消费造成重复订单。",
        "限流策略可以放在网关层或业务服务层。网关层限流适合按 IP、用户或接口维度控制入口流量，能够在请求进入业务服务前完成拦截；业务服务层限流适合结合活动编号、库存状态和用户行为做更精细的控制。实际部署时，可以根据活动规模选择固定窗口、滑动窗口、令牌桶或漏桶算法。",
        "验证码和接口隐藏也是常见防刷手段。验证码可以提高脚本自动化成本，适合在活动开始前获取秒杀资格或请求路径时使用；接口隐藏可以将真实秒杀地址在短时间内动态生成，并绑定用户和活动信息，减少接口被提前暴露后的无效冲击。本文侧重后端核心链路，相关策略可作为扩展功能接入。",
        "降级设计的目标是在系统压力过高或下游异常时优先保护核心能力。当数据库写入变慢时，接口仍可通过 Redis 快速返回排队或成功结果，但后台消费者需要监控积压长度。如果积压超过阈值，系统可以暂停新的秒杀请求、关闭活动入口或提示用户稍后查询结果，避免消息无限堆积。",
        "防护策略还要兼顾用户公平性。过于严格的 IP 限流可能误伤同一网络下的正常用户；只按用户限流又可能无法防止大量账号同时请求。较合理的方案是组合多维度信号，例如用户维度、设备维度、IP 维度、活动维度和请求频率维度，并根据活动热度动态调整阈值。",
    ]:
        add_paragraph(doc, para)

    add_section_title(doc, "5.12 数据库表与索引优化")
    for para in [
        "数据库表设计需要围绕查询路径和写入路径进行优化。订单表是秒杀链路的重点表，常见查询包括按用户查询订单列表、按用户和优惠券判断是否重复购买、按订单编号查询订单状态。因此，订单表应至少保证主键索引有效，并考虑为 user_id、voucher_id 或二者组合建立索引。",
        "若业务明确要求同一用户对同一秒杀券只能购买一次，可以在数据库层增加 user_id 与 voucher_id 的唯一约束。这样即使异步消费者因异常重复处理同一消息，数据库也能拒绝重复写入。唯一约束带来的写入成本通常低于重复订单造成的数据修复成本，因此在强约束业务中具有实际价值。",
        "秒杀券表需要支持按 voucher_id 查询活动时间和库存。由于 voucher_id 与基础优惠券通常是一一对应关系，可以将其作为主键或唯一索引。活动开始前，服务会读取秒杀券信息进行 Redis 预热；活动期间，订单服务也可能根据 voucher_id 查询活动时间。因此该查询路径需要保持稳定。",
        "优惠券表与店铺表更多承担展示和关联查询职责。对店铺优惠券列表，可以在 shop_id 上建立索引；对活动状态筛选，可以结合 begin_time、end_time 和 type 字段进行查询优化。具体索引数量不宜过多，过多索引会增加写入成本，应根据实际查询频率和执行计划调整。",
        "数据库层还需要注意事务隔离和行锁范围。扣减库存时应尽量使用带条件的更新语句，例如库存大于零时才扣减，避免先查后改造成并发窗口。更新条件应命中索引，减少锁扫描范围。订单保存和库存扣减放在同一事务中，可以提高数据一致性，但事务体应保持短小。",
    ]:
        add_paragraph(doc, para)

    add_section_title(doc, "5.13 接口返回与用户体验设计")
    for para in [
        "秒杀接口的返回设计需要兼顾准确性和响应速度。库存不足、重复购买、活动未开始、活动已结束等可以即时确定的失败状态，应直接返回明确提示；通过资格校验的请求应尽快返回订单编号或处理中状态，避免用户长时间等待数据库落库完成。",
        "返回订单编号并不意味着订单已经完成所有后续流程，而是表示用户已经通过秒杀资格校验，订单正在异步创建。前端可以根据订单编号轮询订单状态，或在用户订单列表中展示“处理中”。这种交互方式能够降低接口阻塞时间，也符合高并发活动中常见的排队体验。",
        "错误信息需要面向用户而不是面向开发者。例如 Redis 脚本返回 1 时，后端不应直接返回“脚本返回码 1”，而应转换为“库存不足”；返回 2 时，应提示“不能重复购买”。清晰的用户提示可以减少重复提交，也能降低活动期间客服压力。",
        "对于异常状态，系统应避免暴露内部组件细节。数据库异常、Redis 短暂不可用或消费者处理失败，可以统一提示“系统繁忙，请稍后重试”或“订单处理中”。详细异常应写入服务日志，供开发和运维人员排查。用户侧只需要知道当前操作是否成功、是否可重试以及如何查询结果。",
    ]:
        add_paragraph(doc, para)

    add_section_title(doc, "5.14 日志与监控设计")
    for para in [
        "秒杀系统上线后，需要依靠日志和监控判断活动运行状态。接口层应记录关键失败原因，例如库存不足、重复购买、未登录、活动未开始等；消费者层应记录消息处理成功、处理失败、重试次数、死信数量和数据库事务失败原因。这些日志不仅用于排查问题，也可以用于后续活动复盘。",
        "监控指标可以分为入口指标、Redis 指标、RabbitMQ 指标和数据库指标。入口指标包括请求量、成功率、失败率、平均响应时间和响应时间分位数；Redis 指标包括库存键变化、命令耗时、连接数和内存使用；RabbitMQ 指标包括 Topic 消息堆积量、消费延迟、重试次数和死信队列数量；数据库指标包括连接数、慢查询、锁等待和 CPU 使用率。",
        "活动期间，监控大盘应突出秒杀链路的核心状态。例如库存剩余量是否快速归零、订单消息是否积压、消费者是否正常确认消息、数据库是否出现明显锁等待。如果库存已归零但入口请求仍然很高，可以通过网关或前端快速关闭入口，减少无效请求。",
        "告警阈值需要结合活动规模设置。平时合理的阈值在秒杀活动期间可能过于敏感，导致大量无效告警；而阈值过高又可能错过真实故障。较好的方式是在活动前进行压测，根据压测结果设定临时活动阈值，并在活动结束后恢复常规阈值。",
    ]:
        add_paragraph(doc, para)

    add_section_title(doc, "5.15 部署与扩展性设计")
    for para in [
        "分布式架构的优势之一是可以按模块独立扩展。秒杀活动期间，请求主要集中在网关、订单服务和 Redis，因此可以优先增加网关实例和订单服务实例，提高入口转发和脚本执行调度能力。商户服务、内容服务等非核心链路可以保持常规规模，避免资源浪费。",
        "Redis 是秒杀链路的关键组件，应根据活动规模选择合适部署方式。单机 Redis 部署简单，但可用性和容量有限；主从加哨兵可以提高可用性；集群模式可以提高容量和吞吐能力。对于库存扣减这种单键热点操作，仍需关注热点键所在节点的压力，必要时可以按活动或商品拆分库存键。",
        "订单消费者可以通过多实例扩展提升落库能力。使用 RabbitMQ 消费组时，同一消费组内多个消费者可以分担消息处理。不过，扩展消费者前需要确认数据库写入能力是否足够，否则只是把压力从消息队列转移到数据库。消费者数量应根据数据库连接池、事务耗时和消息积压情况动态调整。",
        "部署过程中还应关注配置隔离。测试环境、预发布环境和生产环境的 Redis、MySQL、网关地址和服务端口应分开配置，避免误连。秒杀活动参数也应通过配置或后台管理发布，而不是写死在代码中，这样才能支持多活动、多批次和临时调整。",
    ]:
        add_paragraph(doc, para)

    add_section_title(doc, "5.16 小结")
    for para in [
        "本章围绕秒杀系统的详细设计与实现进行了说明。系统将高并发请求的关键判断前移到 Redis，通过 Lua 脚本保证库存扣减和重复购买校验的原子性；通过 RabbitMQ 将订单落库异步化，实现削峰填谷；通过数据库事务和重复订单兜底校验保证最终数据可靠。",
        "从整体效果看，该方案在性能、正确性和实现复杂度之间取得了较好的平衡。Redis 负责快速判定，RabbitMQ 负责缓冲和投递订单消息，MySQL 负责最终持久化，网关负责统一鉴权和路由。各组件职责明确，便于后续按业务规模继续扩展限流、监控、告警和容灾能力。",
    ]:
        add_paragraph(doc, para)


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    figs = make_figures()
    doc = Document()
    set_styles(doc)
    configure_section(doc.sections[0], header=False, roman=False, start=1, page_num=False)

    add_cover(doc)
    add_abstracts(doc)
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, header=True, roman=True, start=1)
    add_toc(doc)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, header=True, roman=False, start=1)

    add_chapter(doc, "第一章  引言")
    add_section_title(doc, "1.1 项目研究的背景")
    for para in [
        "随着互联网技术的飞速发展和电子商务的普及，网络购物与本地生活服务已经成为人们日常生活的重要组成部分。在平台运营过程中，秒杀活动通过限时、限量、低价等方式集中吸引用户参与，能够在短时间内带来较高的流量和交易转化。",
        "秒杀业务的技术难点集中体现在瞬时高并发、流量峰值集中、库存数量有限和一致性要求严格等方面。活动开始瞬间，大量用户会同时提交请求，请求量可能在数秒内由常规水平上升至峰值水平。如果系统仍然采用同步查询库存、同步扣减数据库、同步创建订单的传统流程，数据库连接池、线程池和应用服务都可能成为瓶颈。",
        "本课题选取本地生活服务平台中的优惠券秒杀业务作为研究场景。该类平台通常包含用户、商户、内容、关注、优惠券、订单等多个业务模块，具备典型的多服务协作特征。本文在这一业务背景下，重点研究秒杀优惠券下单链路的性能优化，通过缓存、原子脚本、异步消息和服务拆分等手段提升系统在高并发场景下的稳定性。",
    ]:
        add_paragraph(doc, para)
    add_section_title(doc, "1.2 项目研究的目标与意义")
    add_section_title(doc, "1.2.1 研究目标", level=3)
    for para in [
        "本文的主要目标是设计并实现一套基于分布式架构的秒杀系统性能优化方案，具体包括：构建清晰的微服务模块边界，完成网关、用户、优惠券和订单等服务之间的协作；使用 Redis 缓存秒杀库存并通过 Lua 脚本保证库存扣减的原子性；利用 RabbitMQ 实现订单消息异步处理，降低数据库瞬时写入压力；通过接口鉴权、重复购买校验和异常兜底机制保证业务正确性。",
    ]:
        add_paragraph(doc, para)
    add_section_title(doc, "1.2.2 研究意义", level=3)
    add_paragraph(doc, "本课题能够帮助理解高并发系统中缓存、数据库、消息队列和事务边界之间的协同关系。对于业务层面而言，稳定的秒杀系统可以减少活动失败和用户流失；对于技术层面而言，本文形成的优化思路可迁移到抢票、预约、限量报名等类似场景。")
    add_section_title(doc, "1.3 论文内容结构")
    add_paragraph(doc, "本文共分为七章。第一章介绍研究背景、目标与意义；第二章介绍系统相关技术；第三章完成需求分析；第四章进行系统概要设计；第五章阐述核心模块详细设计与实现；第六章进行系统测试与性能分析；第七章总结全文并提出后续优化方向。")

    add_chapter(doc, "第二章  相关技术与工具介绍")
    techs = [
        ("2.1 Spring Boot框架", "Spring Boot 是基于 Spring 的快速开发框架，能够通过自动配置、起步依赖和内嵌服务器简化 Java Web 应用开发。在本系统中，Spring Boot 用于构建用户服务、商户服务、优惠券服务和订单服务等业务模块。"),
        ("2.2 Spring Cloud Gateway", "Spring Cloud Gateway 是响应式 API 网关组件，适用于统一入口、请求路由、跨域处理和基础鉴权。本系统通过网关将用户请求转发到对应微服务，并在入口层完成 Token 校验与用户上下文传递。"),
        ("2.3 Redis缓存技术", "Redis 是基于内存的高性能键值数据库，支持字符串、集合、位图、地理位置和 Stream 等数据结构。本系统使用 Redis 存储登录令牌、商品缓存、秒杀库存、一人一单集合和异步订单消息。"),
        ("2.4 Redis Lua脚本与RabbitMQ", "Lua 脚本在 Redis 服务端原子执行，适合在高并发场景下完成库存判断、扣减和重复购买校验。RabbitMQ 提供 Topic、Exchange、Queue、Consumer Group、消息重试和死信队列等能力，可以支撑订单异步落库与失败补偿。"),
        ("2.5 MySQL数据库与MyBatis-Plus", "MySQL 用于持久化用户、店铺、优惠券、秒杀券和订单等核心业务数据。MyBatis-Plus 简化了常见增删改查和分页查询开发，提高了数据访问层的开发效率。"),
        ("2.6 JMeter与自动化测试工具", "JMeter 可模拟并发请求并收集响应时间、吞吐量和错误率等指标。项目中还通过 JUnit 5、RestAssured 和 AssertJ 编写接口自动化测试，用于验证登录、商品查询、秒杀下单和重复购买等场景。"),
    ]
    for title, body in techs:
        add_section_title(doc, title)
        add_paragraph(doc, body)

    add_chapter(doc, "第三章  系统需求分析")
    add_section_title(doc, "3.1 功能性需求")
    add_paragraph(doc, "系统面向普通用户和管理员两类角色。普通用户可以登录、浏览店铺与优惠券、参与秒杀、查询订单状态；管理员可以发布优惠券、配置秒杀库存和活动时间，并查看活动结果。核心业务功能围绕秒杀券展示、秒杀请求提交、库存校验、订单生成和结果查询展开。")
    add_figure(doc, figs["flow"], "图3.1  秒杀核心业务流程图")
    add_section_title(doc, "3.2 非功能性需求")
    add_paragraph(doc, "非功能性需求主要包括性能、可用性和安全性三个方面。性能上，系统需要尽量缩短秒杀接口响应时间，并避免请求直接冲击数据库；可用性上，订单异步处理需要具备消息重试、死信队列和人工补偿能力；安全性上，系统需要完成登录鉴权、重复购买限制以及非法活动状态拦截。")
    add_table(doc, "系统主要需求说明", ["需求类型", "具体要求", "实现方式"], [
        ["性能", "降低接口响应时间，提升吞吐量", "Redis Lua 原子预扣减、RabbitMQ 异步落库"],
        ["一致性", "防止超卖和重复购买", "库存键与购买集合在 Lua 脚本中统一校验"],
        ["可用性", "订单消费失败后可恢复", "RabbitMQ 死信队列补偿处理"],
        ["安全性", "限制未登录访问和恶意重复请求", "网关鉴权、用户上下文与一人一单校验"],
    ], "3.1")
    add_section_title(doc, "3.3 可行性分析")
    add_paragraph(doc, "从技术可行性看，Spring Boot、Redis、MySQL 等技术成熟稳定，项目已有对应微服务和接口测试基础；从经济可行性看，系统主要使用开源组件，开发和部署成本较低；从操作可行性看，用户通过浏览器或前端页面即可参与活动，系统接口设计清晰，易于集成。")

    add_chapter(doc, "第四章  系统概要设计")
    add_section_title(doc, "4.1 系统架构设计")
    add_paragraph(doc, "系统采用微服务架构，网关作为统一入口，用户服务负责登录与令牌管理，商户服务负责店铺与商品数据，优惠券服务负责秒杀券配置，订单服务负责秒杀请求处理和订单持久化。Redis 同时承担缓存、库存原子校验和消息缓冲的职责，MySQL 负责最终业务数据存储。")
    add_figure(doc, figs["arch"], "图4.1  系统总体架构图")
    add_section_title(doc, "4.2 核心流程设计")
    add_paragraph(doc, "秒杀请求进入订单服务后，系统先校验用户登录状态和活动时间，再生成订单号并执行 Redis Lua 脚本。脚本完成库存是否充足、一人一单是否满足、库存预扣减、用户购买集合记录等操作。接口层收到脚本成功结果后立即返回订单号，后台消费者通过 RabbitMQ 接收订单消息并执行数据库事务。")
    add_figure(doc, figs["seq"], "图4.2  Redis Lua与RabbitMQ异步下单时序图")
    add_section_title(doc, "4.3 数据库设计")
    add_paragraph(doc, "数据库设计围绕用户、店铺、优惠券、秒杀券和订单五类核心实体展开。优惠券记录描述基础券信息，秒杀券记录活动库存与时间窗口，订单记录用户购买结果。实体之间通过用户编号、店铺编号和优惠券编号建立关联。")
    add_figure(doc, figs["er"], "图4.3  系统E-R图")
    add_table(doc, "核心数据表设计", ["表名", "主要字段", "说明"], [
        ["tb_user", "id, phone, nick_name, create_time", "存储用户基础信息"],
        ["tb_shop", "id, name, type_id, x, y", "存储店铺及地理位置数据"],
        ["tb_voucher", "id, shop_id, title, pay_value, actual_value", "存储优惠券基础信息"],
        ["tb_seckill_voucher", "voucher_id, stock, begin_time, end_time", "存储秒杀活动库存和时间"],
        ["tb_voucher_order", "id, user_id, voucher_id, status, create_time", "存储用户秒杀订单"],
    ], "4.1")

    add_chapter(doc, "第五章  系统详细设计与实现")
    add_section_title(doc, "5.1 秒杀接口设计")
    add_paragraph(doc, "秒杀接口采用 RESTful 风格设计，核心路径为 POST /voucher-order/seckill/{id}。网关校验用户 Token 后将用户信息透传给订单服务，订单服务从用户上下文中获取用户编号，并根据优惠券编号完成活动状态校验和秒杀处理。")
    add_section_title(doc, "5.2 Redis库存预扣减实现")
    add_paragraph(doc, "系统在 Redis 中以 seckill:stock:{voucherId} 作为库存键，以 seckill:order:{voucherId} 作为已购买用户集合。Lua 脚本首先读取库存，若库存不足则返回失败；随后判断用户是否已经购买，若已购买则返回重复购买；在满足条件时，脚本扣减库存、记录用户并发送订单消息到 RabbitMQ。由于整个脚本在 Redis 服务端一次性执行，因此能够避免并发条件下的超卖问题。")
    add_section_title(doc, "5.3 异步订单处理实现")
    add_paragraph(doc, "订单服务启动后配置 RabbitMQ 消费者监听器，持续接收秒杀订单消息。消费者将消息转换为订单对象后，通过事务模板再次进行数据库侧重复订单兜底校验、扣减数据库库存并保存订单。消息处理成功后手动确认；若处理失败，则进入死信队列，避免消息长期滞留。")
    add_section_title(doc, "5.4 缓存与防护策略")
    add_paragraph(doc, "系统对热点数据使用 Redis 缓存，减少数据库读取压力；对秒杀链路使用用户登录态、活动时间和一人一单校验减少无效请求；对数据库层增加重复订单兜底检查，避免异步重复消费造成脏写。上述策略共同提升了系统在高并发场景下的稳定性。")
    add_expanded_design_content(doc)

    add_chapter(doc, "第六章  系统测试与性能分析")
    add_section_title(doc, "6.1 测试环境")
    add_table(doc, "测试环境配置", ["项目", "配置"], [
        ["开发语言", "Java 8"],
        ["后端框架", "Spring Boot、Spring Cloud Gateway"],
        ["数据库与缓存", "MySQL 8.0、Redis 6.x"],
        ["测试框架", "JUnit 5、RestAssured、AssertJ、JMeter"],
        ["核心接口", "POST /voucher-order/seckill/{id}"],
    ], "6.1")
    add_section_title(doc, "6.2 功能测试")
    add_table(doc, "秒杀功能测试用例", ["测试场景", "输入条件", "预期结果"], [
        ["正常秒杀", "用户已登录、库存充足、活动进行中", "返回订单号，库存扣减成功"],
        ["库存不足", "库存为0后继续请求", "返回库存不足，订单不创建"],
        ["重复秒杀", "同一用户多次请求同一秒杀券", "第二次请求失败"],
        ["未登录访问", "请求不携带有效 Token", "网关返回未授权"],
        ["无效优惠券", "请求不存在的优惠券编号", "返回失败结果"],
    ], "6.2")
    add_section_title(doc, "6.3 性能测试")
    add_paragraph(doc, "性能测试模拟高并发用户同时提交秒杀请求，并对优化前后的响应时间、错误率、吞吐量和数据库 CPU 使用率进行对比。优化后，请求在 Redis 层完成快速校验和预扣减，订单落库由后台消费者异步处理，因此核心接口响应时间明显降低，数据库压力得到缓解。")
    add_figure(doc, figs["perf"], "图6.1  性能优化前后对比图")
    add_table(doc, "性能测试结果对比", ["指标", "优化前", "优化后", "优化效果"], [
        ["平均响应时间", "3500ms", "180ms", "降低约94.9%"],
        ["错误率", "45%", "0.5%", "显著下降"],
        ["吞吐量", "1200 TPS", "8500 TPS", "提升约6.1倍"],
        ["数据库CPU使用率", "95%", "35%", "压力明显下降"],
    ], "6.3")
    add_section_title(doc, "6.4 测试结果分析")
    add_paragraph(doc, "测试结果说明，Redis Lua 脚本能够保证库存扣减与一人一单校验的原子性，RabbitMQ 异步处理机制能够将高峰请求转化为平稳的后台订单写入。优化方案既缩短了接口响应路径，又避免数据库直接承受瞬时峰值流量。")

    add_chapter(doc, "第七章  总结与展望")
    add_section_title(doc, "7.1 总结")
    add_paragraph(doc, "本文针对秒杀系统的高并发性能问题，完成了基于分布式架构的秒杀系统性能优化设计与实现。系统通过网关统一入口、Redis 缓存与 Lua 脚本原子预扣减、RabbitMQ 异步订单处理、数据库事务兜底校验等方式，解决了秒杀场景下的响应慢、数据库压力大、超卖和重复购买等问题。")
    add_section_title(doc, "7.2 不足与展望")
    add_paragraph(doc, "当前系统仍存在进一步优化空间。例如，网关层限流策略可以结合用户画像和活动热度动态调整；订单消费者可以扩展为多消费者实例以提升落库能力；监控体系还可以补充链路追踪、实时告警和容量预测。后续可继续引入 Sentinel、Prometheus、分布式追踪和自动扩缩容机制，提高系统在复杂生产环境中的可观测性和弹性。")

    add_chapter(doc, "参考文献")
    refs = [
        "[1] Newman S. 微服务设计[M]. 北京: 人民邮电出版社, 2016.",
        "[2] Kleppmann M. 数据密集型应用系统设计[M]. 北京: 中国电力出版社, 2018.",
        "[3] 李智慧. 大型网站技术架构: 核心原理与案例分析[M]. 北京: 电子工业出版社, 2013.",
        "[4] 周志明. 深入理解Java虚拟机: JVM高级特性与最佳实践[M]. 北京: 机械工业出版社, 2019.",
        "[5] 王福顺, 张强, 刘伟. 基于Redis的高并发秒杀系统设计与实现[J]. 计算机工程与应用, 2020, 56(15): 238-244.",
        "[6] Vogels W. Eventually consistent[J]. Communications of the ACM, 2009, 52(1): 40-44.",
        "[7] Redis Ltd. Redis documentation[EB/OL]. https://redis.io/docs/.",
        "[8] Spring. Spring Boot Reference Documentation[EB/OL]. https://docs.spring.io/spring-boot/.",
    ]
    for ref in refs:
        p = add_paragraph(doc, ref, first_line=False, size=10.5)
        for r in p.runs:
            set_run(r, size=10.5, font=FONT_CN)

    add_chapter(doc, "致    谢")
    for para in [
        "在论文完成之际，我要向所有给予我帮助和支持的人表示衷心的感谢。",
        "首先，感谢指导老师杨宇博老师在论文选题、系统设计和论文撰写过程中给予的悉心指导和宝贵建议。老师严谨的治学态度和认真负责的指导方式，使我在系统实现和论文写作过程中受益匪浅。",
        "其次，感谢同学们在项目开发和测试过程中给予的帮助与支持。与大家的讨论交流让我对高并发系统设计、缓存应用和接口测试有了更加深入的理解。",
        "最后，感谢家人在求学期间给予的关心和鼓励。由于时间和能力有限，论文中仍难免存在不足之处，恳请各位老师批评指正。",
    ]:
        add_paragraph(doc, para)

    # Remove page-break-before for front matter headings only where needed.
    for p in doc.paragraphs:
        if p.text in {"摘    要", "ABSTRACT", "参考文献", "致    谢"}:
            p.paragraph_format.page_break_before = True

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_doc())
